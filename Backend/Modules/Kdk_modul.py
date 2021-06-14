from os import *
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime as dt
from Backend.Modules.Users2 import *
from UI.MainMenu import esc_num

now = dt.now().strftime("%Y.%m.%d")

leagues = ['Первенства г.Москвы по футболу среди команд спортивных школ «Клубная Лига» сезона 2021г.',
           'Первенства г.Москвы по футболу среди команд спортивных школ «Первая Лига» сезона 2021г.',
           'Первенства г.Москвы по футболу среди команд спортивных школ «Вторая Лига» сезона 2021г.',
           'Первенства г.Москвы по футболу среди команд спортивных школ «Третья Лига» сезона 2021г.',
           'Первенства г.Москвы по футболу среди команд спортивных школ «Четвертая Лига» сезона 2021г.',
           'Первенства г.Москвы по футболу среди команд спортивных школ «Пятая Лига» сезона 2021г.',
           'Чемпионата г. Москвы по футболу среди мужчин в рамках зонального этапа Всероссийских '
           'спортивных соревнований по футболу «III дивизион» сезона 2021 года',
           'Кубка города Москвы по футболу имени Роберта Фульды среди команд мальчиков 2008 года рождения',
           'Кубка Москвы по футболу среди команд любительских футбольных клубов']
tribunal = ['Дорохин Евгений Владимирович', 'Дроздов Виталий Вячеславович', 'Новиков Иван Юрьевич',
            'Шейхи Фархад Хавасович', 'Авдонченко Наталья Михайловна']

if __name__ == '__main__':
    path_to_db = getcwd()[0:-16] + '\\DB\\db_config.ini'
else:
    path_to_db = getcwd() + '\\DB\\db_config.ini'

home_cwd = getcwd()


def read_db_config(filename=path_to_db, section='mysql'):
    """ Читает конфигурацию Базы данных и возвращает словарь с параметрами
    :param filename: имя конфига
    :param section: секция с данными базы данных
    :return: словарь с параментрами
    """
    parser = ConfigParser()
    parser.read(filename)
    db = {}
    if parser.has_section(section):
        items = parser.items(section)
        for item in items:
            db[item[0]] = item[1]
    else:
        raise Exception('{0} not found in the {1} file'.format(section, filename))

    return db


db_config = read_db_config()

conn = mysql.connector.connect(**db_config)  # соединение с базой данных

cursor = conn.cursor()  # курсор


def connect(func):
    """Подключение в БД """

    @wraps(func)
    def wrapper(*args, **kwargs):
        connec = mysql.connector.connect(**db_config)

        try:
            # print('Соединение с MySQL базой...')
            if not connec.is_connected():
                # print('соединение установлено.')
                print('соединения нет!!!.')
            # else:
            #     print('соединения нет!!!.')

        except Error as error:
            print(error)
        func_result = func(*args, **kwargs)
        if func_result:
            connec.close()
            # print('Соединение закрыто.')
            return func_result
        else:
            connec.close()
            # print('Соединение закрыто.')

    return wrapper


class Notification:

    def __init__(self, case_id, date, time, member, potential_art, league, notif_folder, call=0, match_video=0):
        self.case_id = case_id
        self.date = date
        self.time = time
        self.member = member
        self.potential_art = potential_art
        self.league = league
        self.call = call
        self.match_video = match_video
        self.text = ''
        self.notif_folder = notif_folder
        self.esc_num = esc_num

    def create_notification(self):
        document = Document()

        # Разметка страницы
        sections = document.sections
        for section in sections:
            section.top_margin = 314400

        # Стили
        font_styles = document.styles
        font_head = font_styles.add_style('HeadStyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_head.font
        font_object.size = Pt(11.5)
        font_object.name = 'Times New Roman'

        font_default = font_styles.add_style('DefaultStyle', WD_STYLE_TYPE.CHARACTER)
        font_object_2 = font_default.font
        font_object_2.size = Pt(13)
        font_object_2.name = 'Times New Roman'

        font_main_text = font_styles.add_style('DefaultMainText', WD_STYLE_TYPE.CHARACTER)
        font_object_3 = font_main_text.font
        font_object_3.size = Pt(13)
        font_object_3.name = 'Times New Roman'

        # Загрузка логотипа и шапки
        document.add_picture(f'{home_cwd}\\img/mff_logo.png', width=Inches(1.25))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        head = document.add_paragraph()
        head.add_run('Региональная общественная организация\n'
                     '«Московская федерация футбола»\n'
                     'Контрольно-дисциплинарный комитет', style='HeadStyle').bold = True
        head.alignment = 1

        # номер сообщения исходящего и дата
        numb = document.add_paragraph()
        numb.add_run(f'Исх. №{esc_num} от {now}', style='DefaultStyle')

        # Клубы-адресаты
        club = document.add_paragraph()
        club.add_run(f'{self.member.name}', style='DefaultStyle').bold = True
        club.alignment = 2

        # Основной абзац
        main_text = document.add_paragraph()
        main_text.add_run(f'Контрольно-дисциплинарным комитетом Региональной общественной организации '
                          f'«Московская федерация футбола» (далее – «КДК РОО МФФ») '
                          f'рассматриваются вопросы о несоблюдении требований Регламента '
                          f'«{self.league}» со стороны {self.member.name}, '
                          f'а именно:', style='DefaultStyle')
        main_text.paragraph_format.first_line_indent = Inches(0.25)
        main_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Нарушения
        for art in range(len(self.potential_art)):
            art_text = document.add_paragraph(f'Нарушение, предусмотренное ст.{self.potential_art[art]} '
                                              f'Дисциплинарного регламента () , в соответствии с которым в '
                                              f'протокол матча Клуба, участвующего в Дивизионе «А», могут '
                                              f'включаться не более 5 (пяти) футболистов по карточке Клуба, '
                                              f'участника Чемпионата в дивизионе «Б», но не старше 21 года, '
                                              f'если учредителем этих клубов является одно юридическое лицо.',
                                              style='List Number 2')
            art_text.paragraph_format.left_indent = Inches(0.025)
            art_text.paragraph_format.first_line_indent = Inches(0.25)
            art_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            art_text.style.font.size = Pt(13)
            art_text.style.font.name = 'Times New Roman'

        # Уведомлем о заседании
        if self.call == 0:
            call_text = document.add_paragraph()
            call_text.add_run(f'Уведомляем Вас, что заседание КДК состоится {self.date} в {self.time} по адресу: '
                              f'город Москва, ул. Радио д.12, стр.2, в помещение конференц-зала, '
                              f'без вызова заинтересованных лиц.',
                              style='DefaultStyle')
        else:
            call_text = document.add_paragraph()
            call_text.add_run(f'Уведомляем Вас, что заседание КДК состоится {self.date} в {self.time} по адресу: '
                              f'город Москва, ул. Радио д.12, стр.2, в помещение конференц-зала, '
                              f'без вызова заинтересованных лиц.',
                              style='DefaultStyle')

        # Обязанность по предоставлению информации
        if self.match_video == 0:
            call_text.paragraph_format.first_line_indent = Inches(0.25)
            call_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            notify_text = document.add_paragraph()
            notify_text.add_run(f'Заинтересованным лицам, в срок до 18:00 {self.date}, '
                                f'необходимо предоставить КДК РОО МФФ свои пояснения с обоснованием своей позиции '
                                f'по рассматриваемым дисциплинарным нарушениям со ссылкой на доказательства '
                                f'(документы направлять по адресу email: Zyxel1995@mail.ru).', style='DefaultStyle')
            notify_text.paragraph_format.first_line_indent = Inches(0.25)
            notify_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            call_text.paragraph_format.first_line_indent = Inches(0.25)
            call_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            notify_text = document.add_paragraph()
            notify_text.add_run(f'Заинтересованным лицам, в срок до 18:00 {self.date}, '
                                f'необходимо предоставить КДК РОО МФФ видео матча и свои пояснения с '
                                f'обоснованием своей позиции по рассматриваемым дисциплинарным нарушениям '
                                f'со ссылкой на доказательства '
                                f'(документы направлять по адресу email: Zyxel1995@mail.ru).', style='DefaultStyle')
            notify_text.paragraph_format.first_line_indent = Inches(0.25)
            notify_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Подпись
        signature_doc = document.add_paragraph()
        signature = signature_doc.add_run(style='DefaultStyle')
        signature.add_text(f'Секретарь\nКонтрольно - дисциплинарного\n')
        signature.alignment = 1
        signature.add_picture(f'{home_cwd}\\img/signature.jpg', width=Inches(1.25))
        signature.add_text(f'       	  А.Е.Иванов\n'
                           f'7(905)7630517')

        # Сохранить документ
        document.save(f'Иcx№{self.esc_num}-{self.member.name}.docx')

        # сохраняем текст уведомления
        doc = Document(f'{self.notif_folder}\\Иcx№{self.esc_num}-{self.member.name}.docx')
        self.text = ''.join([p.text for p in doc.paragraphs])

    @connect
    def add_to_base(self):
        query = f'INSERT INTO Notifications_Cases(case_id,notification) VALUES ' \
                f'({self.case_id}, "{self.text}");'
        cursor.execute(query)
        conn.commit()


class Decesion:
    pass


class Case:

    def __init__(self, meeting_id, date, number, time, members_of_meeting, league,
                 potential_art, call=0, match_video=0, notes=''):
        self.meeting_id = meeting_id
        self.case_id = 0
        self.date = date
        self.number = number
        self.time = time
        self.members_of_meeting = members_of_meeting
        self.potential_art = potential_art
        self.league = league
        self.call = call
        self.match_video = match_video
        self.notes = notes
        self.case_files = ''
        self.my_folder = 0
        self.my_notifications = []
        self.creat_folder()
        self.add_in_base()
        self.add_members_in_base()

    def creat_folder(self):
        members = '-'.join([x.shrt_name for x in self.members_of_meeting])
        if self.my_folder == 0:
            self.my_folder = getcwd() + f'\\MFF_Base\\КДК\\Заседание {self.date}\\{members}'
            mkdir(self.my_folder)
        else:
            print('Папка уже создана')

    @connect
    def add_in_base(self):
        if self.notes == '':
            query = f"INSERT INTO Cases(meeting_id,date_meeting,case_number,time_meeting,potential_art) VALUES " \
                    f"({self.meeting_id}, '{self.date}', '{self.number}', '{self.time}', '{self.potential_art}');"
            cursor.execute(query)
            conn.commit()
        else:
            query = f'INSERT INTO Cases(meeting_id,date_meeting,case_number,time_meeting,potential_art,notes) VALUES ' \
                    f'({self.meeting_id}, "{self.date}", "{self.number}", "{self.time}", ' \
                    f'"{self.potential_art}","{self.notes}");'
            cursor.execute(query)
            conn.commit()
        query = f'select id from Cases where case_number="{self.number}";'
        cursor.execute(query)
        number = cursor.fetchone()[0]
        self.case_id = number

    @connect
    def add_members_in_base(self):
        values = ','.join([f'({self.meeting_id},{member.user_club.user_id})' for member in self.members_of_meeting])
        query = f"INSERT INTO ClubMembersOfMeeting(case_id, club_member_id) VALUES {values}"
        cursor.execute(query)
        conn.commit()

    def creat_notification(self):
        chdir(self.my_folder)
        for member in self.members_of_meeting:
            notif = Notification(self.case_id, self.date, self.time, member, self.potential_art,
                                 self.league, self.my_folder, self.call, self.match_video)
            notif.create_notification()
            notif.add_to_base()
            self.my_notifications.append(notif)

    def creat_desicion(self):
        pass

    def send(self):
        pass


class KDK:
    def __init__(self, date, notes=''):
        self.id = 0
        self.date = date
        self.notes = notes
        self.cases = []
        self.my_folder = 0
        self.creat_folder()
        self.add_in_base()

    def creat_folder(self):
        if self.my_folder == 0:
            self.my_folder = getcwd() + f'\\MFF_Base\\КДК\\Заседание {self.date}'
            mkdir(self.my_folder)
        else:
            print('Папка уже создана')

    @connect
    def add_in_base(self):
        if self.notes == '':
            query = f"INSERT INTO KDK(date_meeting) VALUES {self.date};"
            cursor.execute(query)
            conn.commit()
        else:
            query = f'INSERT INTO KDK(date_meeting,notes) VALUES ("{self.date}","{self.notes}");'
            cursor.execute(query)
            conn.commit()
        query = f'select id from KDK where date_meeting="{self.date}";'
        cursor.execute(query)
        number = cursor.fetchone()[0]
        self.id = number

    def creat_case(self, number, time, members_of_meeting, league, potential_art, call=0, match_video=0, notes=''):
        new_case = Case(self.id, self.date, number, time, members_of_meeting, league,
                        potential_art, call, match_video, notes)
        self.cases.append(new_case)

    def agenda(self):
        if self.cases == []:
            print('Дел для рассмотрения нет')
        pass

    def my_agenda(self):
        if self.cases == []:
            print('Дел для рассмотрения нет')
        pass

    def dec_on_site(self):
        if self.cases == []:
            print('Дел для рассмотрения нет')
        pass
